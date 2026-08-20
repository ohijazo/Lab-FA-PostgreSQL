"""Genera la guia de formacio Lab FA - Tutorial pas a pas amb exemples."""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IMAGES_DIR = r"P:\lab-fa\images"

doc = Document()


def add_image(filename, caption=None, width_cm=15):
    """Afegeix una imatge centrada amb peu opcional."""
    path = os.path.join(IMAGES_DIR, filename)
    if not os.path.exists(path):
        print(f"AVIS: imatge no trobada: {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = cp.add_run(f"Figura: {caption}")
        crun.font.size = Pt(9)
        crun.font.italic = True
        crun.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

# -- Estils globals --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)


def add_step(number, title, description):
    """Afegeix un pas numerat amb titol en negreta i descripcio."""
    p = doc.add_paragraph()
    run = p.add_run(f"Pas {number}. {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    doc.add_paragraph(description)


def add_tip(text):
    """Afegeix un consell/nota destacada."""
    p = doc.add_paragraph()
    run = p.add_run("Consell: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
    p.add_run(text)


def add_warning(text):
    """Afegeix un avis destacat."""
    p = doc.add_paragraph()
    run = p.add_run("Atencio: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
    p.add_run(text)


def add_example_box(title, content_lines):
    """Afegeix una caixa d'exemple amb dades concretes."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Light Shading Accent 1"
    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run(f"Exemple: {title}")
    run.bold = True
    for line in content_lines:
        cell.add_paragraph(line)
    doc.add_paragraph()


# ============================================================
# PORTADA
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Lab FA")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Guia de formacio pas a pas")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Tutorial complet amb exemples del flux diari")
run.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Agost 2026")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_page_break()

# ============================================================
# INDEX
# ============================================================
doc.add_heading("Index", level=1)
index_items = [
    "1. Abans de comencar",
    "2. Entrar a l'aplicacio (login)",
    "3. La pagina principal",
    "4. Consultar la llista d'analisis",
    "   4.1. Cercar per text",
    "   4.2. Ordenar columnes",
    "   4.3. Filtrar per data i estat",
    "5. Crear una nova analisi (exemple complet)",
    "6. Consultar el detall d'una analisi",
    "7. Editar una analisi existent",
    "8. Marcar una analisi com a finalitzada",
    "9. Marcar una alerta amb motiu",
    "10. Duplicar una analisi",
    "11. Enviar una analisi per email",
    "12. Exportar dades a Excel",
    "13. Escaner de codi de barres",
    "14. Dashboard i grafics",
    "15. Configurar columnes de la llista",
    "16. [ADMIN] Gestio d'usuaris",
    "17. [ADMIN] Crear un tipus d'analisi nou",
    "18. [ADMIN] Afegir seccions i camps",
    "19. [ADMIN] Camps calculats amb formules",
    "   19.1. Com definir un camp calculat",
    "   19.2. Sintaxi de les formules",
    "   19.3. Exemples reals",
    "   19.4. Quan es fa el calcul",
    "   19.5. Formules encadenades",
    "   19.6. Casos especials i errors",
    "   19.7. Limitacions",
    "   19.8. Bones practiques",
    "20. Preguntes frequents",
    "21. Referencia rapida",
]
for item in index_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if item.startswith("   "):
        p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# ============================================================
# 1. ABANS DE COMENCAR
# ============================================================
doc.add_heading("1. Abans de comencar", level=1)
doc.add_paragraph(
    "Aquesta guia esta pensada per aprendre a fer servir Lab FA de zero, "
    "seguint els mateixos passos que faras en el teu dia a dia al laboratori. "
    "Cada seccio inclou un exemple concret perque puguis reproduir-lo mentre llegeixes."
)
doc.add_paragraph("El que necessites:")
requirements = [
    "Un ordinador amb navegador web (Chrome, Edge, Firefox).",
    "Un compte d'usuari (l'administrador te l'haura creat).",
    "L'adreca de l'aplicacio (per exemple: http://labfa.empresa.local).",
]
for r in requirements:
    doc.add_paragraph(r, style="List Bullet")

add_tip(
    "Recomanem que segueixis aquesta guia amb l'aplicacio oberta i vagis fent cada "
    "pas alhora que el llegeixes. Aprendras molt mes rapid."
)

# ============================================================
# 2. LOGIN
# ============================================================
doc.add_heading("2. Entrar a l'aplicacio (login)", level=1)
doc.add_paragraph(
    "Abans de poder fer res, cal identificar-se. Aixo permet que l'aplicacio sapiga "
    "qui ha creat o modificat cada analisi."
)

add_step(1, "Obre el navegador",
    "Obre Chrome, Edge o Firefox. No cal instal.lar cap programa.")
add_step(2, "Escriu l'adreca de l'aplicacio",
    "A la barra d'adreces del navegador, escriu l'URL que t'ha donat l'administrador "
    "(per exemple: http://labfa.empresa.local) i prem Enter.")
add_step(3, "Introdueix les credencials",
    "Apareixera una pantalla de login demanant el correu electronic i la contrasenya. "
    "Escriu les teves credencials.")
add_step(4, "Clica el boto 'Entrar'",
    "Si les dades son correctes, entraras directament a la pagina principal.")

add_warning(
    "No comparteixis mai la teva contrasenya amb ningu. Si sospites que algu la sap, "
    "canvia-la immediatament des del teu perfil."
)

add_example_box("Login", [
    "URL: http://labfa.empresa.local",
    "Correu: joan.pages@harinera.cat",
    "Contrasenya: (la teva contrasenya personal)",
])

add_image("login.png", caption="Pantalla d'inici de sessio")

# ============================================================
# 3. PAGINA PRINCIPAL
# ============================================================
doc.add_heading("3. La pagina principal", level=1)
doc.add_paragraph(
    "Un cop dins, veuras la pagina principal (Home). Aqui tens una visio general "
    "de tota l'activitat del laboratori."
)
doc.add_paragraph("Que hi trobes:")
home_items = [
    ("Indicadors globals (KPIs)", "A dalt de tot, 4 targetes amb: total d'analisis, "
     "analisis d'aquest any, analisis d'aquest mes, i data de l'ultima analisi."),
    ("Taula de tipus d'analisis", "Al centre, una taula amb tots els tipus disponibles "
     "(per exemple: Blats T1, Bases, Finals). Cada fila mostra el total, les del mes actual "
     "i tres botons: Nou, Llista, Dashboard."),
    ("Boto 'Exportar Excel'", "A dalt a la dreta, per generar un Excel amb dades de multiples tipus."),
    ("Activitat recent", "Al final, una llista plegable amb les ultimes analisis creades."),
    ("Menu superior", "Enllacos a: Home, Escaner, i (si ets admin) Configuracio."),
]
for title, desc in home_items:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

add_image("home.png", caption="Pagina principal (Home) amb KPIs i taula de tipus")

# ============================================================
# 4. CONSULTAR LA LLISTA
# ============================================================
doc.add_heading("4. Consultar la llista d'analisis", level=1)
doc.add_paragraph(
    "La llista es on veuras totes les analisis d'un tipus concret. "
    "Aqui podras cercar, filtrar, ordenar i accedir al detall."
)

add_step(1, "Des de la Home, clica 'Llista' al tipus que t'interessi",
    "Per exemple, al tipus 'Blats T1', clica el boto 'Llista'. T'obrira l'URL /blats_t1.")
add_step(2, "Observa la taula",
    "Veuras una taula compacta amb les columnes principals: data, codi, analista, bareja, etc. "
    "Cada fila es una analisi.")
add_step(3, "Files vermelles = alerta",
    "Si veus una fila en vermell, es que aquella analisi te una alerta activada. "
    "Veurem com fer-ho a la seccio 9.")

add_image("llista_bases.png", caption="Llista d'analisis amb columnes, cerca i filtres")

doc.add_heading("4.1. Cercar per text", level=2)
add_step(1, "Localitza el camp de cerca",
    "A dalt de la taula hi ha un camp de text 'Cercar...'.")
add_step(2, "Escriu el que busques",
    "Pots buscar per qualsevol camp: codi, nom, analista, bareja... La llista es filtra "
    "en temps real mentre escrius.")

add_example_box("Cercar analisis de l'analista Maria", [
    "1. Vas a /blats_t1",
    "2. Al camp cerca escrius: Maria",
    "3. La taula nomes mostra les analisis on apareix 'Maria'",
])

doc.add_heading("4.2. Ordenar columnes", level=2)
add_step(1, "Clica la capcalera de la columna",
    "Per exemple, clica 'Data' per ordenar per data.")
add_step(2, "Torna a clicar per invertir l'ordre",
    "Les fletxes ↑ ↓ indiquen l'ordre actiu.")

doc.add_heading("4.3. Filtrar per data i estat", level=2)
add_step(1, "Filtre de dates",
    "Sobre la taula tens 'Des de' i 'Fins a' per limitar el rang.")
add_step(2, "Filtre d'estat",
    "El toggle 'Tots / Pendent / Finalitzat' permet veure nomes analisis pendents o acabades.")

# ============================================================
# 5. CREAR UNA NOVA ANALISI (EXEMPLE COMPLET)
# ============================================================
doc.add_heading("5. Crear una nova analisi (exemple complet)", level=1)
doc.add_paragraph(
    "Aquesta es la funcio que mes faras servir. Anem a crear una analisi de Blats T1 "
    "amb dades reals com a exemple. Segueix cada pas amb l'aplicacio oberta."
)

add_step(1, "Des de la Home o des de la Llista, clica 'Nou'",
    "T'obrira l'URL /blats_t1/nou amb un formulari buit organitzat per seccions.")
add_step(2, "Omple la seccio 'Identificacio'",
    "Data d'analisi, codi de lot, analista responsable, bareja... "
    "Els camps marcats amb un asterisc son obligatoris.")
add_step(3, "Omple 'Resultats farina'",
    "Introdueix els valors mesurats: humitat, proteina, W, P, L, etc. "
    "Segons vas escrivint, els camps calculats (marcats amb ƒ) es refresquen automaticament.")
add_step(4, "Omple 'Resultats NIR'",
    "Els valors del NIR: cendres, ghumit, absorcio.")
add_step(5, "Omple la resta de seccions",
    "Segons el tipus, tindras seccions com Trituracio, Compressio, Rendiments, Observacions...")
add_step(6, "Verifica els camps calculats",
    "Baixa a les seccions de rendiments. Els camps amb badge 'ƒ' ja s'han omplert sols. "
    "Comprova que els numeros son coherents.")
add_step(7, "Clica 'Guardar'",
    "El boto es a baix del formulari. Si tot es correcte, l'analisi es guarda i vas a la vista de detall.")

add_warning(
    "Si has deixat un camp obligatori buit, apareixera un avis. Els camps marcats en vermell "
    "son els que falten per omplir."
)

add_example_box("Crear una analisi de Blats T1", [
    "SECCIO Identificacio:",
    "  - Data: 14/08/2026",
    "  - Codi: BT1-2026-0234",
    "  - Analista: Maria Puig",
    "  - Bareja: 65% Nacional + 35% Import",
    "",
    "SECCIO Resultats farina:",
    "  - Humitat: 14.2 %",
    "  - Proteina: 11.8 %",
    "  - W: 245",
    "  - P: 78",
    "  - L: 92",
    "",
    "SECCIO Resultats NIR:",
    "  - Cendres: 0.58 %",
    "  - Ghumit: 32.4 %",
    "  - Absorcio: 58.1 %",
    "",
    "SECCIO Rendiment (calculat automaticament):",
    "  - Rendiment farina: 76.2 % (ƒ)",
    "  - Rendiment sego: 22.4 % (ƒ)",
    "  - Rendiment total: 98.6 % (ƒ)",
])

add_tip(
    "Un cop guardada, l'analisi ja apareixera a la llista i comptara als KPIs de la Home."
)

add_image("NouAnàlisisBases.png", caption="Formulari de nova analisi organitzat per seccions")

# ============================================================
# 6. VEURE DETALL
# ============================================================
doc.add_heading("6. Consultar el detall d'una analisi", level=1)
add_step(1, "Des de la llista, clica sobre la fila",
    "O clica el boto 'Detall'. T'obrira l'URL /blats_t1/[id].")
add_step(2, "Revisa les seccions",
    "Totes les dades estan organitzades igual que al formulari.")
add_step(3, "Observa els colors",
    "Els numeros fora de rang apareixen en vermell (per damunt) o blau (per sota) del limit. "
    "Els camps calculats mostren un badge 'ƒ'.")
add_step(4, "Botons d'accio",
    "A dalt del detall tens: Editar, Duplicar, Eliminar, Marcar finalitzat, Marcar alerta, Enviar email.")

# ============================================================
# 7. EDITAR
# ============================================================
doc.add_heading("7. Editar una analisi existent", level=1)
add_step(1, "Des del detall, clica 'Editar'",
    "T'obrira el mateix formulari que 'Nou' pero amb tots els camps ja omplerts.")
add_step(2, "Modifica els camps que calgui",
    "Els camps calculats es refaran automaticament quan canviis els valors d'entrada.")
add_step(3, "Clica 'Guardar'",
    "Els canvis s'apliquen i tornes al detall.")

add_warning(
    "Si un altre usuari esta editant la mateixa analisi, el sistema et bloquejara l'edicio "
    "durant 5 minuts per evitar sobreescriure els seus canvis. Espera't o parla amb ell."
)

add_image("AvisUsuariEditantRegistre.png",
          caption="Avis quan un altre usuari esta editant l'analisi")
add_image("ConflicteRegistreModificat.png",
          caption="Conflicte: el registre ha estat modificat mentre l'editaves")

# ============================================================
# 8. FINALITZAT
# ============================================================
doc.add_heading("8. Marcar una analisi com a finalitzada", level=1)
doc.add_paragraph(
    "Quan una analisi esta completada i revisada, marca-la com a finalitzada perque "
    "tots els companys sapiguen que ja no s'hi ha de tocar."
)
add_step(1, "Ves al detall de l'analisi",
    "O directament a la llista.")
add_step(2, "Clica la casella '✓' o el boto 'Marcar finalitzat'",
    "A la llista tens una columna amb un tick. Clica-la per canviar l'estat.")
add_step(3, "Verifica el canvi",
    "L'analisi apareixera com a 'Finalitzat' i deixara d'apareixer al filtre 'Pendent'.")

# ============================================================
# 9. ALERTA
# ============================================================
doc.add_heading("9. Marcar una alerta amb motiu", level=1)
doc.add_paragraph(
    "Si detectes un problema en una analisi (contaminacio, resultat anomal, error de mesura...), "
    "marca-la amb una alerta perque destaqui visualment i tothom sapiga que cal revisar-la."
)
add_step(1, "Ves al detall de l'analisi",
    "Trobes el boto 'Marcar alerta' entre els botons d'accio.")
add_step(2, "Escriu el motiu",
    "Explica breument que passa. Aquest text el veuran els companys.")
add_step(3, "Guarda",
    "La fila apareixera en vermell a la llista i es veura una icona ⚠ al detall.")

add_example_box("Alerta per humitat elevada", [
    "Analisi: BT1-2026-0234",
    "Motiu: Humitat 14.2% supera el limit de contracte (14.0%).",
    "        Cal revisar amb el proveidor abans de processar el lot.",
])

add_warning(
    "L'alerta es NOMES per problemes reals que requereixen atencio. "
    "No la facis servir com a recordatori o nota personal."
)

# ============================================================
# 10. DUPLICAR
# ============================================================
doc.add_heading("10. Duplicar una analisi", level=1)
doc.add_paragraph(
    "Util quan crees analisis similars sovint (mateix proveidor, mateixa bareja, etc.). "
    "Duplicar copia totes les dades i ho deixa a punt per fer petits canvis."
)
add_step(1, "Des del detall, clica 'Duplicar'",
    "S'obrira el formulari 'Nou' amb totes les dades pre-omplertes.")
add_step(2, "Canvia el que calgui",
    "Data, codi de lot, resultats concrets...")
add_step(3, "Guarda",
    "Es creara com una analisi nova, sense afectar l'original.")

# ============================================================
# 11. ENVIAR PER EMAIL
# ============================================================
doc.add_heading("11. Enviar una analisi per email", level=1)
doc.add_paragraph(
    "Lab FA permet enviar el contingut d'una analisi per correu electronic directament "
    "des del detall, en format visual amb totes les seccions i camps organitzats en una taula. "
    "Util per compartir resultats amb clients, proveidors o companys sense donar-los acces a l'app."
)

add_step(1, "Ves al detall de l'analisi que vols enviar",
    "Des de la llista, clica sobre la fila.")
add_step(2, "Clica el boto 'Enviar email'",
    "Es troba entre els botons d'accio a dalt del detall. IMPORTANT: aquest boto NOMES "
    "apareix si l'administrador del sistema ha configurat el servei d'enviament. "
    "Si no el veus, contacta amb l'admin.")
add_step(3, "Omple el destinatari",
    "Escriu l'adreca de correu de la persona a qui vols enviar l'analisi. "
    "Nomes es pot enviar a UN destinatari a la vegada (no admet CC ni multiples).")
add_step(4, "Revisa l'assumpte",
    "L'assumpte apareix pre-omplert amb el nom del tipus i la referencia de l'analisi. "
    "Pots modificar-lo si vols.")
add_step(5, "Clica 'Enviar'",
    "El sistema construeix un correu amb totes les dades de l'analisi i l'envia. "
    "Apareix un missatge de confirmacio.")

add_example_box("Enviar analisi BT1-2026-0234 al proveidor", [
    "Destinatari: qualitat@proveidor.cat",
    "Assumpte: Blats T1 - BT1-2026-0234",
    "El correu inclou:",
    "  - Text introductori amb el teu nom i la referencia",
    "  - Totes les seccions de l'analisi en taules amb capcaleres",
    "  - Peu amb qui l'ha enviat i quan",
])

doc.add_heading("11.1. Historial d'enviaments", level=2)
doc.add_paragraph(
    "Cada vegada que envies una analisi per email, queda registrat automaticament. "
    "Al detall de l'analisi, sota les seccions de dades, hi ha una zona plegable "
    "'Historial d'enviaments' amb la llista de tots els correus enviats: "
    "destinatari, assumpte, qui el va enviar i quan."
)

add_tip(
    "L'historial es util per demostrar que un resultat es va comunicar a temps, "
    "o per saber quan tornar a enviar-lo si el destinatari no el troba."
)

doc.add_heading("11.2. Com funciona per dins (per a admins)", level=2)
doc.add_paragraph(
    "Lab FA NO envia els correus directament. Delega l'enviament a un flux de "
    "Microsoft Power Automate, que es qui realment envia el correu des del compte "
    "corporatiu. El sistema fa una crida HTTP al webhook del flux amb les dades del correu."
)
doc.add_paragraph("Requisit per activar la funcio:")
doc.add_paragraph(
    "Cal que l'administrador defineixi la variable d'entorn POWER_AUTOMATE_WEBHOOK_URL "
    "al servidor, apuntant al flux corresponent de Power Automate. Sense aquesta "
    "variable, el boto 'Enviar email' no apareix a la interficie.",
    style="List Bullet"
)

add_warning(
    "El sistema considera que l'email s'ha enviat correctament quan Power Automate "
    "respon amb un codi HTTP 2xx. NO garanteix que el correu arribi realment al "
    "destinatari (podria caure a spam, adreca inexistent, etc.). Si un enviament es critic, "
    "confirma amb el destinatari que l'ha rebut."
)

# ============================================================
# 12. EXPORTAR EXCEL
# ============================================================
doc.add_heading("12. Exportar dades a Excel", level=1)
doc.add_paragraph(
    "Lab FA permet exportar les dades a Excel en qualsevol moment. Util per informes, "
    "auditories, o treballar les dades fora."
)

doc.add_heading("12.1. Exportar un sol tipus", level=2)
add_step(1, "Ves a la llista del tipus (per exemple /blats_t1)",
    "")
add_step(2, "Clica 'Exportar Excel'",
    "Apareixera un dialeg amb rang de dates.")
add_step(3, "Tria 'Des de' i 'Fins a'",
    "Per exemple: 01/07/2026 a 31/07/2026.")
add_step(4, "Clica 'Descarregar'",
    "El navegador et descarregara un fitxer .xlsx amb totes les analisis del rang.")

doc.add_heading("12.2. Exportar multiples tipus", level=2)
add_step(1, "Des de la Home, clica 'Exportar Excel'",
    "")
add_step(2, "Selecciona els tipus que vols",
    "Pots marcar-ne mes d'un. Cada tipus sera una pestanya al fitxer Excel.")
add_step(3, "Tria el rang de dates i descarrega",
    "El fitxer resultant conte una pestanya per tipus.")

add_image("ExportarExcel.png", caption="Dialeg d'exportacio a Excel")

# ============================================================
# 13. ESCANER
# ============================================================
doc.add_heading("13. Escaner de codi de barres", level=1)
doc.add_paragraph(
    "Si tens un lector de codi de barres, pots buscar analisis molt rapidament."
)
add_step(1, "Clica 'Escaner' al menu superior",
    "S'obrira una pantalla d'escaneig.")
add_step(2, "Escaneja el codi",
    "El lector el llegeix i el sistema busca l'analisi automaticament.")
add_step(3, "Si existeix, vas al detall",
    "Si no existeix, es proposa crear-ne una nova amb aquell codi.")

# ============================================================
# 14. DASHBOARD
# ============================================================
doc.add_heading("14. Dashboard i grafics", level=1)
doc.add_paragraph(
    "Cada tipus d'analisi te el seu propi dashboard amb estadistiques visuals."
)
add_step(1, "Des de la Home o la llista, clica 'Dashboard'",
    "T'obrira /dashboard/[tipus] amb grafics i indicadors.")
add_step(2, "Explora els indicadors per seccio",
    "Cada camp numeric mostra mitjana, minim i maxim del periode.")
add_step(3, "Filtra per rang de dates",
    "Els grafics es refan segons el filtre.")
add_step(4, "Filtra per categoria",
    "Pots comparar mitjanes per proveidor, tipus de blat, etc.")

# ============================================================
# 15. COLUMNES
# ============================================================
doc.add_heading("15. Configurar columnes de la llista", level=1)
doc.add_paragraph(
    "Cada usuari pot triar quines columnes veu a la llista de cada tipus."
)
add_step(1, "A la llista, clica 'Configurar columnes'",
    "S'obrira un panel amb totes les columnes disponibles.")
add_step(2, "Marca/desmarca les que vols",
    "Nomes es mostraran les marcades.")
add_step(3, "Arrossega per reordenar",
    "Drag-drop per canviar l'ordre.")
add_step(4, "Guarda",
    "La configuracio es guarda per al teu usuari.")

add_image("ColumnesDeLaLLista.png", caption="Configuracio de columnes visibles a la llista")

# ============================================================
# 16. ADMIN: USUARIS
# ============================================================
doc.add_heading("16. [ADMIN] Gestio d'usuaris", level=1)
doc.add_paragraph(
    "Nomes els usuaris amb rol admin poden accedir a aquestes funcions."
)
add_step(1, "Ves a Configuracio > Usuaris",
    "")
add_step(2, "Clica 'Nou usuari'",
    "Omple: correu, nom, rol (viewer/user/admin), contrasenya temporal.")
add_step(3, "Guarda",
    "L'usuari ja pot fer login.")

doc.add_paragraph("Rols disponibles:")
roles = [
    ("viewer", "Nomes lectura. No pot crear ni editar."),
    ("user", "Pot crear, editar, marcar finalitzat i alertes."),
    ("admin", "Tot l'anterior + gestio de tipus, seccions, camps i usuaris."),
]
for role, desc in roles:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{role}: ")
    run.bold = True
    p.add_run(desc)

add_warning(
    "No esborris un usuari que hagi creat analisis. Millor desactiva'l."
)

add_image("Gestio_usuaris.png", caption="Panell de gestio d'usuaris (nomes admins)")
add_image("NouUsuari.png", caption="Formulari de creacio de nou usuari")

# ============================================================
# 17. ADMIN: TIPUS NOU
# ============================================================
doc.add_heading("17. [ADMIN] Crear un tipus d'analisi nou", level=1)
doc.add_paragraph(
    "Lab FA es totalment dinamic: pots crear tipus nous sense tocar codi."
)
add_step(1, "Ves a /admin/tipus (Configuracio > Tipus d'analisis)",
    "")
add_step(2, "Clica 'Nou tipus'",
    "Omple: nom (visible), slug (usat a l'URL), descripcio.")
add_step(3, "Guarda",
    "El tipus queda creat pero encara sense seccions ni camps.")

add_example_box("Crear tipus 'Farines especials'", [
    "Nom: Farines especials",
    "Slug: farines_especials",
    "Descripcio: Analisis de farines destinades a productes especifics (integrals, sense gluten...)",
])

add_warning(
    "El slug NO es pot canviar despres. Escull-lo be. Ha de ser curt, sense espais, "
    "sense accents (per exemple: farines_esp, no 'Farines Especials!')."
)

add_image("CopnfiguracióTipusAnàlisi.png", caption="Panell d'administracio de tipus d'analisi")

# ============================================================
# 18. ADMIN: SECCIONS I CAMPS
# ============================================================
doc.add_heading("18. [ADMIN] Afegir seccions i camps", level=1)

doc.add_heading("18.1. Crear seccions", level=2)
add_step(1, "Des de la llista de tipus, clica 'Seccions' al tipus",
    "T'obre /admin/tipus/[id]/seccions.")
add_step(2, "Clica 'Nova seccio'",
    "Introdueix el titol (per exemple 'Dades generals') i l'ordre.")
add_step(3, "Guarda i repeteix per cada seccio necessaria",
    "Recomanem entre 3 i 6 seccions per tipus.")

add_image("ConfiguracióSeccions.png", caption="Gestio de seccions d'un tipus d'analisi")

doc.add_heading("18.2. Crear camps", level=2)
add_step(1, "Des de la seccio, clica 'Camps'",
    "T'obre /admin/seccions/[id]/camps.")
add_step(2, "Clica 'Nou camp'",
    "Omple: nom intern (sense espais), etiqueta (visible), tipus (text, number, date, select), obligatori.")
add_step(3, "Si es tipus 'select', afegeix les opcions",
    "Una per linia.")
add_step(4, "Guarda",
    "El camp ja apareixera al formulari del tipus.")

add_example_box("Crear camp 'Humitat' a la seccio 'Resultats'", [
    "Nom intern: humitat",
    "Etiqueta: Humitat (%)",
    "Tipus: number",
    "Obligatori: Si",
    "Limit min (alerta): 12.0",
    "Limit max (alerta): 15.0",
])

add_image("ConfiguracióCamps.png", caption="Llistat de camps d'una seccio")
add_image("ConfiguracióNouCamp.png", caption="Formulari per definir un nou camp")

# ============================================================
# 19. ADMIN: FORMULES
# ============================================================
doc.add_heading("19. [ADMIN] Camps calculats amb formules", level=1)
doc.add_paragraph(
    "Els camps calculats s'omplen automaticament a partir d'altres camps del formulari. "
    "Son ideals per calcular rendiments, sumes, mitjanes, diferencies i qualsevol valor "
    "derivat d'altres mesures. L'usuari no els omple: apareixen sols mentre escriu."
)

doc.add_heading("19.1. Com definir un camp calculat", level=2)
add_step(1, "Crea un camp nou (o edita'n un d'existent)",
    "Ves a Configuracio > Seccions > Camps. Recomanem tipus 'number', pero funciona amb qualsevol.")
add_step(2, "Al camp 'Formula', escriu l'expressio",
    "Fes servir els NOMS INTERNS (name) dels altres camps, no les etiquetes visibles. "
    "Per exemple, si un camp te etiqueta 'Humitat (%)' pero name intern 'humitat', "
    "a la formula escrius 'humitat'.")
add_step(3, "Guarda el camp",
    "Al formulari de crear/editar analisi, aquest camp apareixera amb un badge ƒ al costat "
    "i es refrescara automaticament quan canvii qualsevol camp del qual depengui.")

add_tip(
    "El badge ƒ que apareix al costat de l'etiqueta del camp mostra la formula "
    "quan hi passes el ratol per sobre (tooltip). Util per veure d'on ve el valor."
)

doc.add_heading("19.2. Sintaxi de les formules", level=2)
doc.add_paragraph("Operadors aritmetics:")
ops = [
    ("+", "Suma"),
    ("-", "Resta (i canvi de signe: -x)"),
    ("*", "Multiplicacio"),
    ("/", "Divisio"),
    ("( )", "Parentesis per agrupar"),
]
for op, desc in ops:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{op}  ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph("Funcions disponibles:")
funcs = [
    ("min(a, b, ...)", "Valor minim d'una llista de valors."),
    ("max(a, b, ...)", "Valor maxim d'una llista de valors."),
    ("abs(x)", "Valor absolut (sempre positiu)."),
    ("round(x)", "Arrodonir x a l'enter mes proper."),
    ("round(x, n)", "Arrodonir x a n decimals."),
]
for op, desc in funcs:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{op}: ")
    run.bold = True
    p.add_run(desc)

add_warning(
    "Els decimals s'escriuen SEMPRE amb PUNT a la formula: 14.2, no 14,2. "
    "Els usuaris si que poden escriure amb coma als camps del formulari (el sistema "
    "ho converteix), pero a la definicio de la formula ha de ser punt."
)

doc.add_heading("19.3. Exemples reals", level=2)
add_example_box("Formules tipiques", [
    "SUMA DE COMPONENTS:",
    "  rendiment_total = trit_farina + trit_sego",
    "",
    "DIFERENCIA (sempre positiva):",
    "  diferencia = abs(pes_inicial - pes_final)",
    "",
    "MITJANA ARRODONIDA A 2 DECIMALS:",
    "  mitjana = round((v1 + v2 + v3) / 3, 2)",
    "",
    "PERCENTATGE:",
    "  percent_humit = round(humit / pes_total * 100, 1)",
    "",
    "MAXIM D'UNA SERIE DE MESURES:",
    "  maxim_w = max(w_28min, w_1h, w_2h)",
    "",
    "COMBINACIO AMB PARENTESIS:",
    "  index = round((a + b) * (c - d) / 100, 3)",
])

doc.add_heading("19.4. Quan es fa el calcul", level=2)
doc.add_paragraph(
    "El calcul es fa NOMES al navegador de l'usuari, en TEMPS REAL, mentre esta "
    "omplint el formulari. Cada vegada que l'usuari toca qualsevol camp, totes les "
    "formules es recalculen instantaniament."
)
doc.add_paragraph(
    "El servidor NO recalcula res: guarda el valor tal com arriba del navegador. "
    "Aixo vol dir que si canvies una formula despres de crear analisis, les analisis "
    "antigues NO es refan automaticament: mantenen el valor amb que es van guardar."
)

add_tip(
    "Si vols recalcular una analisi antiga despres de canviar la formula, obre-la en "
    "mode 'Editar', toca qualsevol camp perque es dispari el recalcul, i torna a guardar."
)

doc.add_heading("19.5. Formules encadenades", level=2)
doc.add_paragraph(
    "Una formula pot fer servir el resultat d'un altre camp calculat. El sistema "
    "resol les dependencies iterativament (fins a 20 passades) fins que tots els "
    "valors estan estabilitzats."
)

add_example_box("Formules encadenades", [
    "Camp 'subtotal'    (formula):  a + b",
    "Camp 'iva'         (formula):  subtotal * 0.21",
    "Camp 'total'       (formula):  subtotal + iva",
    "",
    "Al canviar 'a' o 'b', el sistema calcula primer 'subtotal',",
    "despres 'iva', i finalment 'total'. Tot automatic.",
])

doc.add_heading("19.6. Casos especials i errors", level=2)

cases = [
    ("Camp d'entrada buit",
     "Si un camp referenciat esta buit o no es numeric, el resultat de la formula queda "
     "en BLANC (no es 0). Aixo evita mostrar valors falsos abans d'omplir totes les dades."),
    ("Divisio per zero",
     "Si la formula divideix per un valor 0, el resultat queda en blanc."),
    ("Error de sintaxi",
     "Si la formula te un error (parentesi sense tancar, nom mal escrit, funcio inventada...), "
     "el camp queda en blanc SILENCIOSAMENT. No es mostra cap missatge d'error a l'usuari, "
     "per aixo cal PROVAR la formula despres de guardar-la."),
    ("Camps de tipus checkbox",
     "Si referencies un camp booleana (checkbox), es tracta com 1 si esta marcat i 0 si no."),
    ("Arrodoniment automatic",
     "Els resultats es mostren amb un maxim de 4 decimals. Si vols menys, fes servir round()."),
]
for title, desc in cases:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

add_warning(
    "IMPORTANT: sempre que crees o modifiques una formula, fes una prova real: "
    "ves al formulari de nova analisi, omple els camps d'entrada amb valors coneguts "
    "i verifica que el resultat es el que esperaves. Un error d'ortografia al nom "
    "d'un camp deixa el resultat buit sense avisar."
)

doc.add_heading("19.7. Limitacions", level=2)
doc.add_paragraph("El sistema de formules es DELIBERADAMENT SIMPLE i segur. Coses que NO permet:")
limits = [
    ("Sense condicionals",
     "No hi ha 'if/else' ni operador ternari. Si necessites un calcul condicional "
     "(p.ex. 'si a > 10, torna b, si no torna c'), no es pot fer amb aquest sistema."),
    ("Sense comparacions",
     "No hi ha operadors <, >, =, !=."),
    ("Sense funcions estadistiques sobre llistes",
     "No hi ha sum(), avg(), count(). Si vols sumar 5 camps, cal escriure "
     "a + b + c + d + e explicitament."),
    ("Sense potencies ni arrels",
     "No hi ha ^ ni pow() ni sqrt(). Per elevar al quadrat: x * x."),
    ("Sense funcions de text o data",
     "Nomes valors numerics. Els camps de tipus text o data no es poden fer servir "
     "com a origen de calcul (retornen buit)."),
]
for title, desc in limits:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

add_tip(
    "Aquestes limitacions son intencionades: mantenen les formules simples de llegir, "
    "impossibles de trencar accidentalment, i segures (el sistema NO fa servir eval() "
    "ni cap execucio de codi arbitrari)."
)

doc.add_heading("19.8. Bones practiques", level=2)
practices = [
    "Fes servir noms de camp CURTS i CLARS (p.ex. 'humit', no 'humitat_de_la_mostra_1').",
    "Documenta la formula al camp de descripcio si es complexa.",
    "Prova sempre la formula amb dades reals abans de donar-la per bona.",
    "Si una formula queda molt llarga, considera dividir-la en dos camps calculats encadenats.",
    "Fes servir round() al final per controlar la precisio de la sortida.",
    "Evita crear dependencies circulars (camp A depen de B, i B depen de A): el resultat sera imprevisible.",
]
for pr in practices:
    doc.add_paragraph(pr, style="List Bullet")

# ============================================================
# 20. PREGUNTES FREQUENTS
# ============================================================
doc.add_heading("20. Preguntes frequents", level=1)

faqs = [
    ("He guardat una analisi amb dades incorrectes. Que faig?",
     "Ves al detall, clica 'Editar', corregeix i guarda de nou. Les dades antigues se sobreescriuen."),
    ("Puc esborrar una analisi?",
     "Si, des del detall amb el boto 'Eliminar'. Es demanara confirmacio. "
     "Es una accio no reversible - fes-la nomes si estas segur."),
    ("El formulari no em deixa guardar",
     "Segurament falta un camp obligatori. Busca els que estan marcats en vermell o "
     "tenen un asterisc."),
    ("Un camp calculat no dona el resultat esperat",
     "Verifica que tots els camps d'entrada estan omplerts. Si algun camp esta buit, "
     "el resultat queda en blanc (no es 0). Comprova tambe que a la formula fas servir "
     "els noms interns dels camps, no les etiquetes visibles."),
    ("Vull veure analisis de fa mesos i no les trobo",
     "Comprova que no tens un filtre de dates actiu. Neteja els filtres i torna a buscar."),
    ("No em deixa entrar (login incorrecte)",
     "Verifica correu i contrasenya. Si has oblidat la contrasenya, contacta amb l'administrador."),
    ("La fila d'una analisi apareix vermella",
     "Significa que te una alerta activada. Obre el detall per veure el motiu."),
    ("Un altre usuari esta editant una analisi que necessito",
     "El sistema bloqueja edicio simultania durant 5 minuts. Espera't o parla amb l'usuari."),
]
for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(f"P: {q}")
    run.bold = True
    doc.add_paragraph(f"R: {a}")

# ============================================================
# 21. REFERENCIA RAPIDA
# ============================================================
doc.add_heading("21. Referencia rapida", level=1)

doc.add_heading("URLs principals", level=2)
urls = [
    ("/", "Home - visio general"),
    ("/login", "Pantalla de login"),
    ("/:tipus", "Llista d'analisis d'un tipus (ex: /blats_t1)"),
    ("/:tipus/nou", "Formulari per crear nova analisi"),
    ("/:tipus/:id", "Detall d'una analisi"),
    ("/:tipus/:id/editar", "Editar analisi"),
    ("/dashboard/:tipus", "Dashboard amb grafics"),
    ("/admin/tipus", "[Admin] Gestio de tipus"),
    ("/admin/tipus/:id/seccions", "[Admin] Seccions d'un tipus"),
    ("/admin/seccions/:id/camps", "[Admin] Camps d'una seccio"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "URL"
hdr[1].text = "Que fa"
for h in hdr:
    for p in h.paragraphs:
        p.runs[0].font.bold = True
for url, desc in urls:
    row = table.add_row().cells
    row[0].text = url
    row[1].text = desc

doc.add_paragraph()

doc.add_heading("Botons clau", level=2)
buttons = [
    ("Nou", "Crear una nova analisi"),
    ("Editar", "Modificar una analisi existent"),
    ("Detall", "Veure totes les dades d'una analisi"),
    ("Duplicar", "Copiar una analisi per crear-ne una similar"),
    ("Eliminar", "Esborrar (no reversible)"),
    ("✓ Finalitzat", "Marcar com a completada"),
    ("⚠ Alerta", "Marcar amb un problema (fila vermella)"),
    ("Exportar Excel", "Descarregar dades en .xlsx"),
    ("Dashboard", "Veure grafics i estadistiques"),
    ("Escaner", "Buscar per codi de barres"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Boto"
hdr[1].text = "Accio"
for h in hdr:
    for p in h.paragraphs:
        p.runs[0].font.bold = True
for btn, desc in buttons:
    row = table.add_row().cells
    row[0].text = btn
    row[1].text = desc

doc.add_paragraph()

doc.add_heading("Icones", level=2)
icons = [
    ("ƒ", "Camp calculat automaticament amb formula"),
    ("✓", "Analisi finalitzada"),
    ("⚠", "Analisi amb alerta"),
    ("↑ ↓", "Ordre actiu a la columna"),
    ("*", "Camp obligatori"),
]
for icon, desc in icons:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{icon}  ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()

# ============================================================
# ============================================================
# PART II - NOVETATS (Agost 2026)
# ============================================================
# ============================================================
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PART II - NOVETATS")
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Actualitzacio Agost 2026")
run.font.size = Pt(14)
run.font.italic = True
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
doc.add_paragraph()
doc.add_paragraph(
    "Aquesta segona part documenta les millores i noves funcionalitats "
    "afegides al sistema durant l'agost de 2026. Inclou nous rols, "
    "noves pantalles, millores d'usabilitat i sistemes de valoracio "
    "que abans no existien."
)
doc.add_paragraph()

# ============================================================
# 22. Estat APTE / NO APTE
# ============================================================
doc.add_heading("22. Estat APTE / NO APTE d'una analisi", level=1)
doc.add_paragraph(
    "Cada analisi te ara una valoracio de qualitat que indica si el "
    "producte analitzat es apte per al seu us. Aquesta valoracio la fa "
    "el laborant despres d'introduir les dades i complementa els estats "
    "existents (Finalitzat, Alerta)."
)
doc.add_paragraph("Els tres estats possibles son:")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Pendent (per defecte): ").bold = True
p.add_run("l'analisi encara no s'ha valorat. Badge amb outline gris.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("APTE: ").bold = True
p.add_run("el producte es correcte i pot ser utilitzat. Badge blau.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("NO APTE: ").bold = True
p.add_run("el producte no compleix els requisits. Badge vermell.")

add_image("apte_badges_detall.png", caption="Badges Finalitzat, Alerta i APTE a la capsalera del detall")

doc.add_heading("22.1. Com marcar una analisi", level=2)
doc.add_paragraph("Tens dues formes de canviar l'estat APTE:")
add_step(1, "Badge clicable a la capcalera del detall",
         "Al detall d'una analisi, al costat del titol veuras els tres badges: "
         "Finalitzat/Pendent, Alerta, i APTE/NO APTE/Pendent. Cliqueu al badge "
         "APTE per anar cicilant: Pendent -> APTE -> NO APTE -> Pendent.")
add_step(2, "Bloc 'Valoracio' al final del detall",
         "Al final de la pagina de detall hi ha un panel destacat amb tres "
         "botons grans: Finalitzat, Alerta i APTE. Ideal per marcar rapidament "
         "sense haver de fer scroll amunt fins als badges.")

add_tip(
    "El nou bloc 'Valoracio' al final del detall es especialment util quan "
    "acabes d'introduir dades: no cal fer scroll amunt per marcar els estats."
)

add_image("apte_bloc_valoracio.png", caption="Bloc 'Valoracio' amb 3 botons grans al final del detall")

doc.add_heading("22.2. Visualitzacio", level=2)
doc.add_paragraph(
    "L'estat APTE es visualitza a diversos llocs:"
)
p = doc.add_paragraph(style="List Bullet")
p.add_run("Detall: ").bold = True
p.add_run("badge blau/vermell/pendent al costat del titol i al bloc 'Valoracio' final.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Llistat d'analisis: ").bold = True
p.add_run("nova columna d'icona (check blau si APTE, X vermell si NO APTE, buit si pendent).")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Pantalla de Recepcio: ").bold = True
p.add_run("badge gran a cada anàlisi (veure capitol 24).")

add_warning(
    "Les analisis existents abans d'aquesta feature queden com a 'Pendent'. "
    "S'han de marcar manualment un cop revisades."
)

add_image("apte_columna_llista.png", caption="Nova columna d'estat APTE (check blau) i NO APTE (X vermella) al llistat")

# ============================================================
# 23. Rol Recepcio (crear l'usuari)
# ============================================================
doc.add_heading("23. [ADMIN] Rol Recepcio", level=1)
doc.add_paragraph(
    "S'ha creat un nou rol anomenat 'Recepcio' pensat especificament per al "
    "personal de recepcio de camions (el noi que rep els camions al pati)."
)
doc.add_paragraph("Aquest rol:")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Nomes te acces a una unica pantalla: /recepcio")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Si intenta anar a qualsevol altra URL, es redirigeix automaticament a /recepcio")
p = doc.add_paragraph(style="List Bullet")
p.add_run("No pot crear, editar ni eliminar analisis")
p = doc.add_paragraph(style="List Bullet")
p.add_run("No veu el menu de configuracio ni cap altra opcio")
p = doc.add_paragraph(style="List Bullet")
p.add_run("El seu login redirigeix automaticament a la pantalla de recepcio")

doc.add_heading("23.1. Crear un usuari amb rol Recepcio", level=2)
add_step(1, "Anar a la gestio d'usuaris",
         "Configuracio -> Gestio d'usuaris. Nomes admins hi tenen acces.")
add_step(2, "Nou usuari",
         "Botó 'Nou usuari' a dalt a la dreta. Empleneu email, nom complet i contrasenya.")
add_step(3, "Rol: Recepcio",
         "Al selector de rol, tria 'Recepcio' (l'ultima opcio). Els altres rols son: "
         "Editor, Administrador, Lectura.")
add_step(4, "Crear",
         "Clic a 'Crear'. Ja pots facilitar les credencials al noi de recepcio.")

add_image("rol_recepcio_selector.png", caption="Selector de rol amb la nova opcio 'Recepcio'")

add_tip(
    "Configura el navegador del tablet o ordinador de la caseta de recepcio "
    "perque obri directament l'aplicacio (com a pinned tab). "
    "Un cop logat, el sistema es queda a /recepcio automaticament."
)

# ============================================================
# 24. Pantalla de Recepcio
# ============================================================
doc.add_heading("24. Pantalla de Recepcio", level=1)
doc.add_paragraph(
    "Pantalla dedicada per al personal de recepcio, dissenyada per un us rapid "
    "amb tablet o ordinador dedicat. Mostra en temps real les analisis fetes "
    "avui perque el recepcionista sapiga si un cami que arriba te el producte "
    "APTE o NO APTE."
)

doc.add_heading("24.1. Que veu el recepcionista", level=2)
doc.add_paragraph("La pantalla es divideix en:")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Capçalera: ").bold = True
p.add_run("titol 'Recepcio', indicador 'En directe' (puntet verd pulsant), boto per activar avisos sonors.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Barra de cerca: ").bold = True
p.add_run("busca per codi, proveïdor o num de tiquet.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Tabs d'estat: ").bold = True
p.add_run("[Tots] [Pendent] [APTE] [NO APTE] amb comptadors per cada estat.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Filtres secundaris: ").bold = True
p.add_run("desplegable de tipus + botons de rang d'hora (Tots / Ultimes 4h / Ultima hora).")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Llista de files: ").bold = True
p.add_run("una fila per analisi amb codi, proveïdor, tiquet, tipus, hora i badge d'estat.")

add_image("recepcio_pantalla.png", caption="Pantalla de Recepcio: header amb 'En directe', tabs d'estat, filtres i llista de files")

doc.add_heading("24.2. Refresc en temps real (SSE)", level=2)
doc.add_paragraph(
    "La pantalla es refresca automaticament quan hi ha canvis (nova analisi, "
    "canvi d'estat APTE) sense haver de fer F5. Retard: aproximadament 2 segons."
)
doc.add_paragraph(
    "L'indicador 'puntet verd pulsant' a la capsalera confirma que hi ha connexio "
    "activa. Si es posa taronja i parpelleja, vol dir que s'ha perdut la connexio "
    "(el sistema reintenta reconnectar automaticament)."
)

doc.add_heading("24.3. Notificacions sonores", level=2)
doc.add_paragraph(
    "Al costat de l'indicador 'En directe' hi ha un boto 'Activar avisos' (icona "
    "de campana). Un cop activat, la pantalla emet un so quan:"
)
p = doc.add_paragraph(style="List Bullet")
p.add_run("Arriba una analisi nova: 1 to neutre")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Una analisi passa a APTE: 2 tons pujant (positiu)")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Una analisi passa a NO APTE: 3 tons baixant (alerta)")

add_warning(
    "Els sons sempre funcionen, però les notificacions del sistema operatiu "
    "(les que apareixen a la cantonada de la pantalla) requereixen que l'app "
    "s'accedeixi per HTTPS. A HTTP nomes tens sons. Si veus un badge groc '!' "
    "al costat del boto 'Avisos actius' vol dir que nomes tens sons."
)

doc.add_heading("24.4. Detalls d'una fila (camps d'Identificacio)", level=2)
doc.add_paragraph(
    "Cada fila te un chevron a l'esquerra. Cliquem-la per desplegar els camps "
    "de la seccio 'Identificacio' del tipus (data, proveïdor, sitja, varietat, "
    "matricula camio, matricula remolc, etc.). Serveix perque el recepcionista "
    "confirmi que la fila correspon efectivament al cami que te al davant."
)

doc.add_heading("24.5. Colors dels estats", level=2)
p = doc.add_paragraph(style="List Bullet")
p.add_run("APTE: ").bold = True
p.add_run("border-left blau + badge blau amb check. Deixar descarregar el cami.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("NO APTE: ").bold = True
p.add_run("border-left vermell + badge vermell amb X + fons rosa clar. NO descarregar.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Pendent: ").bold = True
p.add_run("border-left dashed + badge outline. El laboratori encara no ha valorat.")

# ============================================================
# 25. Rangs condicionals per valor
# ============================================================
doc.add_heading("25. [ADMIN] Rangs condicionals per valor", level=1)
doc.add_paragraph(
    "Els rangs valids d'un camp numeric (min i max) poden dependre del valor "
    "d'un altre camp de l'analisi. Exemple classic: la humitat, gluten i "
    "proteïna esperada d'un blat canvia segons el 'tipus de blat' (F1 - Gran "
    "Força, F2 - Força, Panificable, etc.). Aquesta funcionalitat s'anomena "
    "'rangs condicionals per un camp controlador'."
)
doc.add_paragraph(
    "L'agost de 2026 s'ha reforcat molt aquesta funcionalitat amb un wizard "
    "guiat, autosave, vista comparativa, bulk apply, importacio entre tipus i "
    "feedback visual del rang esperat al formulari."
)

add_image("rangs_card_seccions.png",
          caption="Card 'Rangs condicionals per valor' a la pantalla de seccions amb el wizard de 2 passos")

doc.add_heading("25.1. Concepte", level=2)
doc.add_paragraph("Necessites tres elements:")
p = doc.add_paragraph(style="List Number")
p.add_run("Un camp de tipus 'Llista' (select) ").bold = True
p.add_run("que fara de 'controlador' del tipus. Per exemple: 'tipus_blat' amb opcions F1, F2, Panificable...")
p = doc.add_paragraph(style="List Number")
p.add_run("Camps numerics ").bold = True
p.add_run("(humitat, gluten, proteïna, W, P/L, etc.) que tindran rangs diferents segons el valor del controlador.")
p = doc.add_paragraph(style="List Number")
p.add_run("Rangs configurats ").bold = True
p.add_run("des de la pantalla de rangs per cada combinacio valor-controlador × camp × min/max.")

add_example_box("Blats amb tipus_blat com controlador", [
    "F1 - Gran Força: humitat 12-14, gluten 30-35, proteïna 12-14",
    "F2 - Força: humitat 11-15, gluten 28-33, proteïna 11-13",
    "Panificable: humitat 10-15, gluten 26-32, proteïna 10-12",
    "Quan el laborant tria tipus_blat=F1, els camps es pinten d'alerta si surten dels rangs de F1.",
])

doc.add_heading("25.2. Pas 1: crear el camp controlador", level=2)
add_step(1, "Afegir un camp de tipus 'Llista'",
         "Configuracio -> Tipus -> Blats -> Seccions -> [Camps d'una seccio] -> Nou camp. "
         "Al tab 'General', tria tipus 'Llista (combobox)'.")
add_step(2, "Definir les opcions",
         "Un cop el tipus es 'select', apareix el tab 'Opcions'. Afegeix cada valor possible "
         "(F1 - Gran Força, F2 - Força, Panificable...) un per un o via import massiu.")
add_step(3, "Guardar",
         "Boto 'Crear' (o 'Desar canvis' si estas editant).")

doc.add_heading("25.3. Pas 2: seleccionar el camp controlador", level=2)
add_step(1, "Anar a Configuracio -> Tipus -> Blats -> Seccions",
         "Al final de la pagina veuras una nova card destacada 'Rangs condicionals per valor'.")
add_step(2, "Card en estat 'Pas 1: selecciona el camp controlador'",
         "La card mostra un wizard de 2 passos. Pas 1 actiu: dropdown 'Camp controlador' "
         "amb els camps de tipus llista disponibles.")
add_step(3, "Tria i desar",
         "Selecciona 'tipus_blat' (o el nom del camp) al dropdown. Clic 'Desar'. "
         "Un cop desat, apareix el pas 2 amb CTA 'Configurar rangs'.")

doc.add_heading("25.4. Pas 3: wizard de configuracio de rangs", level=2)
add_image("rangs_wizard_valors.png",
          caption="Wizard de rangs: tabs per cada valor del controlador, toolbar amb accions i camps min/max amb autosave")
doc.add_paragraph(
    "Al clicar 'Configurar rangs' vas a /admin/tipus/[id]/rangs. La pantalla te un "
    "wizard focalitzat en un valor cada cop:"
)
p = doc.add_paragraph(style="List Bullet")
p.add_run("Tabs a dalt: ").bold = True
p.add_run("un per cada valor del controlador (F1, F2, Panificable...) amb comptador N/total. "
          "Verd quan tots els camps estan configurats per aquell valor.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Toolbar del valor actiu: ").bold = True
p.add_run("nom, progres, i botons d'accions ('Copiar de...', 'Aplicar a...', 'Netejar').")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Camps agrupats per seccio: ").bold = True
p.add_run("cada camp te dos inputs (min i max) que es desen automaticament al fer blur o Enter.")

add_tip(
    "Autosave silencios: cada input es desa al canviar de camp. Veuras un spinner "
    "durant el desat i despres un tic verd d'1.5 segons. No cal cap boto 'Desar' - "
    "es impossible perdre canvis."
)

doc.add_heading("25.5. Copiar rangs entre valors", level=2)
doc.add_paragraph(
    "Si els rangs de dos valors del controlador son similars, pots copiar-los per "
    "estalviar temps."
)
add_step(1, "Configurar el valor origen",
         "Empleneu tots els camps min/max per F1 - Gran Força (per exemple).")
add_step(2, "Anar al valor destí",
         "Clica al tab F2 - Força.")
add_step(3, "Copiar de...",
         "Al toolbar, dropdown 'Copiar de...' -> tria F1 -> confirma. Els rangs "
         "de F1 es copien a F2 (i ja pots ajustar-los).")

add_image("rangs_copiar_de.png",
          caption="Dropdown 'Copiar de…' obert amb els valors del controlador que ja tenen rangs")

doc.add_heading("25.6. Bulk apply (aplicar a diversos valors)", level=2)
doc.add_paragraph(
    "Si vols que 3 o mes valors comparteixin els mateixos rangs, pots aplicar-los "
    "d'un cop."
)
add_step(1, "Configurar el valor font",
         "Empleneu els rangs per F1 - Gran Força.")
add_step(2, "Boto 'Aplicar a...'",
         "Al toolbar, clic a 'Aplicar a...'. Obre un modal amb checkboxes de tots "
         "els altres valors.")
add_step(3, "Seleccionar destins",
         "Marca F2, F3 i Panificable. Cada checkbox mostra si el valor ja té rangs. "
         "Botons 'Seleccionar tots' / 'Desseleccionar' per rapidesa.")
add_step(4, "Aplicar",
         "Clic 'Aplicar a N valors'. Confirmació d'estil 'F1 sera copiat a: F2, F3, Panificable. "
         "Els valors existents seran sobreescrits. Continuar?'.")

add_image("rangs_aplicar_a.png",
          caption="Modal 'Aplicar a…' amb checkboxes de valors destí i botons de selecció ràpida")

doc.add_heading("25.7. Vista comparativa (matriu de resum)", level=2)
doc.add_paragraph(
    "Al header hi ha un boto 'Vista comparativa' que obre una modal ampla amb "
    "una matriu compacta noms-lectura de tots els valors × camps × rangs. "
    "Colors:"
)
p = doc.add_paragraph(style="List Bullet")
p.add_run("Verd: ").bold = True
p.add_run("rang condicional configurat per aquesta combinacio.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Groc (cursiva): ").bold = True
p.add_run("sense rang condicional, però el camp té rangs estàtics que s'usaran com fallback.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Gris: ").bold = True
p.add_run("cap rang (ni condicional ni estàtic).")

add_tip(
    "La vista comparativa es ideal per revisar la configuracio sencera d'un cop "
    "d'ull i detectar buits (celles grises) o combinacions que faltaven."
)

add_image("rangs_vista_comparativa.png",
          caption="Vista comparativa: matriu valors × camps amb codis de color (verd/groc/gris)")

doc.add_heading("25.8. Importar rangs d'un altre tipus", level=2)
doc.add_paragraph(
    "Si tens dos tipus similars (ex: blats i blats-especials) i vols reutilitzar "
    "la configuracio, pots importar-la."
)
add_step(1, "Boto 'Importar d'un altre tipus'",
         "Al header de la pantalla de rangs. Obre un modal amb els tipus que ja tenen "
         "rangs configurats.")
add_step(2, "Seleccionar tipus font",
         "Tria el tipus del qual vols copiar (ex: 'blats').")
add_step(3, "Revisar previsualitzacio",
         "El sistema mostra: quants camps coincideixen per nom intern, quins valors "
         "del controlador son comuns i quins rangs es filtraran.")
add_step(4, "Importar",
         "Clic 'Importar'. Els rangs del tipus font es copien als camps coincidents del "
         "tipus destí. Els valors del controlador font que no existeixen al destí "
         "s'ignoren automaticament.")

add_warning(
    "L'importacio sobreescriu els rangs existents dels camps coincidents. "
    "Fes servir amb precaucio si ja tenies configuracions manuals que vols preservar."
)

add_image("rangs_importar_tipus.png",
          caption="Modal 'Importar d'un altre tipus' amb previsualització de camps i valors comuns")

doc.add_heading("25.9. Al formulari: rang esperat visible", level=2)
doc.add_paragraph(
    "Un cop configurats els rangs, al formulari d'introduccio d'analisi el laborant "
    "veura sota cada camp numeric el rang esperat:"
)
add_example_box("Formulari nou analisi - Blats amb tipus_blat=F1", [
    "Humitat %      [ 13.2 ]",
    "               Esperat: 12-14           <- text mut gris",
    "",
    "Proteina %     [ 11.5 ]                 <- input vermell",
    "               Esperat: 12-16 (fora)   <- text vermell",
    "",
    "Gluten %       [ 32   ]",
    "               Esperat: 30-35",
])
doc.add_paragraph(
    "El text 'Esperat: X-Y' es re-calcula automaticament quan el laborant canvia "
    "el valor del controlador (tipus_blat). Aixo li permet saber els rangs sense "
    "haver de recordar-los ni consultar-los."
)
doc.add_paragraph(
    "Casos especials:"
)
p = doc.add_paragraph(style="List Bullet")
p.add_run("Nomes minim: 'Esperat: ≥ 12'")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Nomes maxim: 'Esperat: ≤ 14'")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Controlador buit + camp amb rangs condicionals: hint 'Requereix «Tipus de blat»'")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Camp sense cap rang configurat: no es mostra res")

doc.add_paragraph(
    "Al detall de l'analisi, els valors amb alerta pintada tambe mostren el rang "
    "entre parentesis: '13.5 (esperat: 12-14)'."
)

add_image("rangs_formulari_esperat.png",
          caption="Formulari: hint 'Esperat: X – Y' sota cada camp numèric; camp fora del rang pintat de vermell")

# ============================================================
# 26. Millores d'usabilitat
# ============================================================
doc.add_heading("26. Millores d'usabilitat", level=1)
doc.add_paragraph(
    "Petites millores generals a l'aplicacio que fan l'us diari mes agil."
)

doc.add_heading("26.1. Cerca live sense Intro", level=2)
doc.add_paragraph(
    "Al llistat d'analisis (/blats, /bases, etc.) ja no cal premer Intro despres "
    "d'escriure al camp de cerca. Al cap de 300ms d'aturar-te de teclejar, la "
    "llista es filtra sola. Intro segueix funcionant per forçar cerca instantania."
)

doc.add_heading("26.2. Keyboard shortcuts globals", level=2)
doc.add_paragraph("Nous atalls de teclat que funcionen fora d'inputs:")
shortcuts = [
    ("/", "Focus al camp de cerca (Llistat, Home, Recepcio)"),
    ("n", "Nou analisi (des del llistat)"),
    ("e", "Editar analisi (des del detall)"),
    ("Escape", "Tancar modals oberts"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Light Shading Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Tecla"
hdr[1].text = "Accio"
for h in hdr:
    for pp in h.paragraphs:
        pp.runs[0].font.bold = True
for tecla, desc in shortcuts:
    row = table.add_row().cells
    row[0].text = tecla
    row[1].text = desc
doc.add_paragraph()

add_tip(
    "Els shortcuts s'ignoren mentre estas escrivint dins d'un input, textarea "
    "o select (excepte Escape que sempre funciona). Aixi no interfereixen amb "
    "el que estas escrivint."
)

doc.add_heading("26.3. Sticky 'Desar' als formularis", level=2)
doc.add_paragraph(
    "Als formularis llargs (com el de blats amb 7 seccions), el boto 'Desar' es "
    "queda enganxat al peu de pantalla mentre scroleges. No has de baixar tot al "
    "fons per desar."
)

doc.add_heading("26.4. Empty states amb accions rapides", level=2)
doc.add_paragraph(
    "Quan una llista esta buida (Home sense tipus creats, Dashboard sense dades, "
    "Admin sense camps, etc.), en lloc d'un espai buit apareix un card destacada "
    "amb icona + descripcio + boto d'accio directa ('Crear el primer tipus', "
    "'Nova analisi'...). Aixi es mes facil saber que fer a continuacio."
)

doc.add_heading("26.5. Confirmacions d'eliminacio mes segures", level=2)
doc.add_paragraph(
    "Els modals de confirmacio d'eliminacio ara mostren el nom/codi de l'element "
    "en NEGRETA per evitar equivocacions:"
)
add_example_box("Confirmar eliminacio", [
    "Abans: 'Segur que vols eliminar aquesta analisi?'",
    "Ara:   'Segur que vols eliminar l'analisi 3044?'  (3044 en negreta)",
])

doc.add_heading("26.6. Refresc en temps real (recepcio)", level=2)
doc.add_paragraph(
    "La pantalla de recepcio utilitza Server-Sent Events per mostrar els canvis "
    "en temps real. Quan un laborant crea una analisi o canvia l'estat APTE, el "
    "recepcionista ho veu al cap de 2 segons aproximadament sense fer F5. "
    "Vegis capitol 24 per detalls."
)

doc.add_heading("26.7. Feedback en camps numerics", level=2)
doc.add_paragraph(
    "Els camps numerics del formulari mostren el rang esperat sota l'input "
    "('Esperat: 12-14'), i canvien de color d'alerta si el valor introduit surt "
    "del rang. El laborant sap el valor esperat abans d'introduir-lo. "
    "Vegis capitol 25.9 per detalls."
)

doc.add_heading("26.8. Responsive tablet-first", level=2)
doc.add_paragraph(
    "L'aplicacio s'adapta a pantalles de tablet (iPad portrait 768px i landscape "
    "1024px). El menu, els botons i els formularis es reorganitzen per no quedar "
    "tallats. Aixo permet fer servir l'app comodament des d'un tablet a la caseta "
    "de recepcio o al laboratori."
)

doc.add_heading("26.9. Notificacions al recepcionista", level=2)
doc.add_paragraph(
    "El recepcionista pot activar sons (i notificacions del sistema si l'app "
    "esta en HTTPS) per saber quan arriba una analisi nova o quan passa a APTE/NO APTE. "
    "Vegis capitol 24.3 per detalls."
)

# Peu final
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Lab FA - Guia de formacio")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Per qualsevol dubte, contacta amb l'administrador del sistema.")
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

# -- Guardar --
output_path = r"P:\lab-fa\Lab_FA_Guia_Formacio.docx"
doc.save(output_path)
print(f"Document generat: {output_path}")
